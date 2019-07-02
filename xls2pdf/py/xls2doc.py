# coding=utf-8
# @Time    : 2019/6/19 9:35
# @Author  : Leau
# @File    : xls2doc.py
#
try:
    import os
    os.chdir(r'./')
    from win32com.client import Dispatch,constants,gencache
except Exception as e:
    print(str(e))
#   input('Please press enter key to exit0 ...')
#%%
def word2Pdf(wordFile,pdfFile):
    try:
        print('Word Excel Export ...')
        word=gencache.EnsureDispatch('Word.Application')
        doc=word.Documents.Open(wordFile,ReadOnly=1)
#        input('Open ...')
        doc.ExportAsFixedFormat(pdfFile,
                constants.wdExportFormatPDF,
                Item=constants.wdExportDocumentWithMarkup,
                CreateBookmarks=constants.wdExportCreateHeadingBookmarks)
        print("pdfFile:",pdfFile)
        word.Quit(constants.wdDoNotSaveChanges)
    except Exception as e:
        print(str(e))
        #让程序停在这里等待回车键退出
        #input('Please press enter key to exit1 ...')
        #%%
def excel2Pdf(excelFile,pdfFile):
    try:
        print('Begin Excel Export ...')
        xlApp=Dispatch('Excel.Application')
        books=xlApp.Workbooks.Open(excelFile)
        books.ExportAsFixedFormat(0,pdfFile)
        print("pdfFile:",pdfFile)
        xlApp.Quit()
    except Exception as e:
        print(str(e))
        #让程序停在这里等待回车键退出
#       input('Please press enter key to exit1 ...')
#%%
def ExportWord():
    wordFiles=[fn for fn in os.listdir('.') if fn.endswith(('.doc','.docx')) ]
    for wordFile in wordFiles:
        wordFile=os.path.abspath(wordFile)
        index=wordFile.rindex('.')
        pdfFile=wordFile[:index]+'.pdf'
        word2Pdf(wordFile,pdfFile)
#%%
def ExportExcel():
    try:
        excelFiles=[fn for fn in os.listdir('.') if fn.endswith(('.xls','.xlsx')) ]
        for excelFile in excelFiles:
            #print("excelFiles:",excelFiles)
            excelFile=os.path.abspath(excelFile)
            index=excelFile.rindex('.')
            pdfFile=excelFile[:index]+'.pdf'
            excel2Pdf(excelFile,pdfFile)
    except Exception as e:
        print(str(e))
# 让程序停在这里等待回车键退出
#         input('Please press enter key to exit3 ...')

# %%

print("开始导出。。。")
# ExportWord()
# ExportExcel()
print("导出完成。。。")
from pdfminer.pdfparser import PDFParser, PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.layout import LAParams
from pdfminer.converter import PDFPageAggregator
from pdfminer.pdfinterp import PDFTextExtractionNotAllowed
from docx import Document
document = Document()
import warnings
warnings.filterwarnings("ignore")
import os
# 添加你要转换的文件地址
file_name = os.open(r'./pdf/山海关水泉市场提升改造工程.pdf', os.O_RDWR)
def main():
    fn = open(file_name, 'rb')
    parser = PDFParser(fn)
    doc = PDFDocument()
    parser.set_document(doc)
    doc.set_parser(parser)
    resource = PDFResourceManager()
    laparams = LAParams()
    device = PDFPageAggregator(resource, laparams=laparams)
    interpreter = PDFPageInterpreter(resource, device)
    for i in doc.get_pages():
        interpreter.process_page(i)
        layout = device.get_result()
        for out in layout:
            if hasattr(out, "get_text"):
                content = out.get_text().replace(u'\xa0', u' ')
                document.add_paragraph(
                    content, style='ListBullet'
                )
            document.save('a1' + '.docx')
    print('处理完成')

if __name__ == '__main__':
    main()